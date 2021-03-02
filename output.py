# output.py
# project:sudoku
# littzhch
# 20200716
# 20200717


from docx import Document
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
import json


def _json_format(source_str):
    """
    格式化json字符串
    """
    state = 0
    digit_count = 0
    result = ''

    for char in source_str:
        if state == 0:
            result += char
        elif state == 1:
            result += char
            if char == ',':
                result += "\n\n\n\t\t"
        elif state == 2:
            if char.isdigit():
                result += char
                digit_count += 1
            elif char == ',':
                result += char
                if not digit_count % 27:
                    result += "\n\n\t\t "
                elif not digit_count % 9:
                    result += "\n\t\t "
                elif not digit_count % 3:
                    result += "    "
            elif char == ']':
                result += char
                digit_count = 0

        if char == '[':
            state += 1
        elif char == ']':
            state -= 1

    return result


def _set_cell_border(cell, **kwargs):
    """
    This function comes from
    https://www.jianshu.com/p/9ad7db7825ba
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def _add_single_square(docu, sdk):
    """
    添加单个数独
    :param docu: docx.Document对象
    :param sdk: sudoku.SudokuSquare对象
    :return: 无
    """
    t = docu.add_table(rows=9, cols=9, style="Table Grid")  # 创建表格
    t.style.font.size = Pt(22)  # 设置字体大小
    t.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格居中对齐
    for row in t.rows:
        row.height = Cm(1)  # 设置行高

    idx = 0

    for a in range(0, 9):
        for b in range(0, 9):

            cell = t.cell(a, b)

            num = sdk.raw_info[idx]  # 添加数字
            if num == 0:
                text = ''
            else:
                text = str(num)
            cell.text = text
            idx += 1

            cell.width = Cm(1)  # 设置列宽
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER               # 设置单元格居中
            cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

            if a == 0:                                                                    # 设置边框
                _set_cell_border(cell, top={"sz": 22, "val": "single"})
            if a == 8:
                _set_cell_border(cell, bottom={"sz": 22, "val": "single"})
            if a == 3:
                _set_cell_border(cell, top={"sz": 16, "val": "single"})
            if a == 6:
                _set_cell_border(cell, top={"sz": 16, "val": "single"})
            if b == 0:
                _set_cell_border(cell, left={"sz": 22, "val": "single"})
            if b == 8:
                _set_cell_border(cell, right={"sz": 22, "val": "single"})
            if b == 3:
                _set_cell_border(cell, left={"sz": 16, "val": "single"})
            if b == 6:
                _set_cell_border(cell, left={"sz": 16, "val": "single"})


def create_docx(file_name, sdks, answers):
    """
    创建并保存数独和答案的docx文件
    :param file_name: 数独文件名
    :param sdks: 数独的sudoku.SudokuSquare对象列表
    :param answers: 答案的sudoku.SudokuSquare对象列表
    :return: 无
    """
    s_file = Document()
    a_file = Document()
    idx = 0
    pagebreak = False
    while idx <= len(sdks) - 1:
        sdks[idx].idx_update()
        s_file.add_paragraph("题目" + str(idx + 1) + "(" + str(len(sdks[idx].zero_idx)) + "空)：")
        a_file.add_paragraph("题目" + str(idx + 1) + "答案：")
        _add_single_square(s_file, sdks[idx])
        _add_single_square(a_file, answers[idx])
        idx += 1
        if pagebreak:
            if idx != len(sdks):
                s_file.add_page_break()
                a_file.add_page_break()
            pagebreak = False
        else:
            s_file.add_paragraph("")
            a_file.add_paragraph("")
            pagebreak = True
    s_file.save(file_name + ".docx")
    a_file.save(file_name + "答案.docx")


def create_json(file_name, sdks, answers):
    """
    创建并保存数独和答案的json文件
    :param file_name: 数独文件名
    :param sdks: 数独的sudoku.SudokuSquare对象列表
    :param answers: 答案的sudoku.SudokuSquare对象列表
    :return: 无
    """
    rst_dict = {
        "problems" : [],
        "answers"  : []
        }

    for idx in range(len(sdks)):
        rst_dict["problems"].append(sdks[idx].raw_info)
        rst_dict["answers"].append(answers[idx].raw_info)

    with open(file_name + '.json', "w") as f:
        f.write(_json_format(json.dumps(rst_dict, indent=4)))
